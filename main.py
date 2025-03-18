"""
This file contains FileChat-RAG program implemented using Python.
FileChat-RAG is a simple Retrieval-Augmented Generation (RAG) system that allows users to ask questions about the contents of various file formats.
It extracts text from PDFs, JSON, text files(.txt, .md), document files(.docx, .odt), and code files(.py, .cpp, .java, .c, .js, .ts, .html, .csharp, .sh), then enables interactive conversations using an LLM powered by Ollama.
"""
import argparse
import fitz
from docx import Document
from odf.opendocument import load
from odf.text import P
import os
import json
from langchain_ollama import OllamaLLM
from langchain_community.vectorstores import FAISS
from langchain_ollama import OllamaEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import ConversationalRetrievalChain

parser = argparse.ArgumentParser(prog='RAG: LLM for doing conversation from data present in a file.')
parser.add_argument("--path", "-p", default="./cd data.pdf")
args = parser.parse_args()
FILE_PATH = args.path

def extract_text_from_txt(file_path):
    """Extract text from a Text (.txt, .md) file."""
    with open(file_path, "r", encoding="utf-8") as file:
        return file.read()

def extract_text_from_docx(file_path):
    """Extract text from a Word Document (.docx), including paragraphs, tables, headers, and footers."""
    doc = Document(file_path)
    text_content = []
    # Extract paragraphs
    for para in doc.paragraphs:
        text_content.append(para.text)
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            text_content.append("\t".join(cell.text for cell in row.cells))
    # Extract headers & footers
    for section in doc.sections:
        if section.header:
            text_content.append("[Header] " + section.header.paragraphs[0].text)
        if section.footer:
            text_content.append("[Footer] " + section.footer.paragraphs[0].text)
    return "\n".join(text_content)

def extract_text_from_odt(file_path):
    """Extract text from an OpenDocument Text (.odt) file."""
    doc = load(file_path)
    text_content = []
    for element in doc.getElementsByType(P): 
        if element.childNodes:
            text_content.append("".join(node.data for node in element.childNodes if node.nodeType == node.TEXT_NODE))
    return "\n".join(text_content)

def extract_text_from_pdf(file_path):
    """Extract text from a PDF (.pdf) file."""
    doc = fitz.open(file_path)
    text = "\n".join(page.get_text("text") for page in doc)
    return text

def extract_text_from_json(file_path):
    """Extract text from a JSON (.json) file."""
    with open(file_path, "r", encoding="utf-8") as file:
        data = json.load(file)
        return json.dumps(data, indent=2)  # Convert JSON to formatted text

def extract_text_from_code(file_path):
    """Extract text from code files (.py, .cpp, .java, .c, .js, .ts, .html, .csharp, .sh)"""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
        return file.read()  # Treat as plain text

def extract_text(file_path):
    """Detect file type and extract text accordingly."""
    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()
    if file_extension in [".txt", ".md"]:
        return extract_text_from_txt(file_path)
    elif file_extension == ".docx":
        return extract_text_from_docx(file_path)
    elif file_extension == ".odt":
        return extract_text_from_odt(file_path)
    elif file_extension == ".pdf":
        return extract_text_from_pdf(file_path)
    elif file_extension == ".json":
        return extract_text_from_json(file_path)
    elif file_extension in [".py", ".cpp", ".java", ".c", ".js", ".ts", ".html", ".csharp", ".sh"]:
        return extract_text_from_code(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")

def get_file_retriever(filename):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    text_data = extract_text(filename)
    docs = text_splitter.create_documents([text_data])
    # vector Store
    embedding_model = OllamaEmbeddings(model="all-minilm")
    vectorstore = FAISS.from_documents(docs, embedding_model)
    # retriever
    retriever = vectorstore.as_retriever()
    return retriever

# Initialize LLM and retriever
llm = OllamaLLM(model="gemma3:1b", base_url="http://localhost:11434")
retriever = get_file_retriever(FILE_PATH)
conversation_chain = ConversationalRetrievalChain.from_llm(llm, retriever=retriever)

def launch_tui():
    print("Active data repository: "+FILE_PATH)
    chat_history = []  # Store previous interactions
    while True:
        query = input("Ask me anything! (Type /quit to exit) >>> ")
        if query == "/quit":
            print("Bye.")
            break
        response = conversation_chain.invoke({"question": query, "chat_history": chat_history})
        # Store conversation history
        chat_history.append((query, response["answer"]))
        print(">>> "+response["answer"])

launch_tui()